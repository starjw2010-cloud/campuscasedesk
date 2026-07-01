#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""make_case_docs.py — 케이스 채널 스레드 공유용 도메인 문서 생성 (PDF + Word) → demo_docs/
실제 양식·체크리스트·안내문 느낌. 전부 synthetic_demo. 사용: python3 make_case_docs.py
"""
from pathlib import Path
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, ListFlowable, ListItem
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).resolve().parent / "demo_docs"; OUT.mkdir(exist_ok=True)
pdfmetrics.registerFont(TTFont("KR", "/System/Library/Fonts/Supplemental/AppleGothic.ttf"))
NAVY = colors.HexColor("#1B3A5B"); GOLD = colors.HexColor("#C8A13A")
ss = getSampleStyleSheet()
def S(sz, col=colors.black, al=0, sp=6):
    return ParagraphStyle(f"s{sz}{al}{sp}", parent=ss["Normal"], fontName="KR", fontSize=sz, leading=sz*1.5, textColor=col, alignment=al, spaceAfter=sp)
TI=S(17,NAVY,1,4); SU=S(10,colors.grey,1,12); H=S(12.5,NAVY,0,5); B=S(10,colors.black,0,4); SM=S(8.5,colors.grey,0,2)
def hdr(t,s): return [Paragraph("한빛대학교 CampusFlow · synthetic_demo",SM),Spacer(1,3),Paragraph(t,TI),Paragraph(s,SU)]
def bl(xs): return ListFlowable([ListItem(Paragraph(x,B),leftIndent=10) for x in xs],bulletType="bullet",start="•",bulletColor=GOLD)
def tbl(data,w):
    t=Table(data,colWidths=w); t.setStyle(TableStyle([("FONTNAME",(0,0),(-1,-1),"KR"),("FONTSIZE",(0,0),(-1,-1),9),
      ("BACKGROUND",(0,0),(-1,0),NAVY),("TEXTCOLOR",(0,0),(-1,0),colors.white),
      ("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white,colors.HexColor("#EEF2F7")]),
      ("GRID",(0,0),(-1,-1),0.4,colors.HexColor("#CCCCCC")),("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5)])); return t
def pdf(fn, flow):
    SimpleDocTemplate(str(OUT/fn),pagesize=A4,topMargin=18*mm,bottomMargin=16*mm,leftMargin=18*mm,rightMargin=18*mm).build(flow); print("  ✓ PDF",fn)

def docx_base(title, sub):
    d=Document(); n=d.styles["Normal"]; n.font.name="AppleGothic"; n.font.size=Pt(10.5)
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("한빛대학교 CampusFlow · synthetic_demo"); r.font.size=Pt(8.5); r.font.color.rgb=RGBColor(0x88,0x88,0x88)
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(title); r.bold=True; r.font.size=Pt(16); r.font.color.rgb=RGBColor(0x1B,0x3A,0x5B)
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(sub); r.font.size=Pt(9.5); r.font.color.rgb=RGBColor(0x66,0x66,0x66)
    d.add_paragraph(""); return d
def dh(d,t):
    p=d.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(12.5); r.font.color.rgb=RGBColor(0x1B,0x3A,0x5B)

# 1 현장실습 협약서(PDF)
f=hdr("현장실습 표준 협약서 (양식)","산학협력팀 · 현장실습 운영")
f+=[Paragraph("1. 협약 당사자",H),bl(["대학: 한빛대학교 경영대학","실습기관: (기관명)","학생: (성명/학번 — 마스킹)"]),
    Paragraph("2. 실습 기간·시간",H),tbl([["항목","내용"],["실습기간","2026-09-01 ~ 2026-12-19"],["주당 시간","최대 40시간"],["보험","산재·상해보험 가입 필수"]],[60*mm,114*mm]),
    Paragraph("3. 필수 제출 서류",H),bl(["실습협약서","보험가입증명서","실습일지(주별)","출근부","기관평가서(종료 시)"]),
    Spacer(1,8),Paragraph("※ 중도 포기 시 진단서·기관확인서·지도교수확인서 제출 후 학점 처리 규정에 따름.",SM)]
pdf("현장실습_표준협약서.pdf",f)

# 2 산재보험 안내(PDF)
f=hdr("현장실습 산재·상해보험 가입 안내","산학협력팀")
f+=[Paragraph("가입 대상·시점",H),bl(["전 현장실습 학생 의무 가입","실습 시작 전 가입 완료 — 미가입 시 실습 불가","보험가입증명서를 협약서와 함께 제출"]),
    Paragraph("사고 발생 시 절차",H),bl(["즉시 기관 담당자·지도교수에게 통보","사고경위서·진단서 제출","보험 처리 및 학점 처리 협의"])]
pdf("산재보험_가입안내.pdf",f)

# 3 장학 중복수혜 심사기준(PDF)
f=hdr("장학금 중복 수혜 심사 기준","학생지원팀 · 장학")
f+=[Paragraph("중복 수혜 판정",H),bl(["교내·국가·외부 장학 동시 수혜 시 등록금 초과분 조정","성적우수+가계곤란 중복은 규정상 1개로 제한","수혜내역확인서로 중복 여부 확인"]),
    Paragraph("심사 제출 서류",H),tbl([["서류","비고"],["장학신청서","필수"],["수혜내역확인서","중복 확인용"],["성적증명서","자격 확인"]],[60*mm,114*mm]),
    Spacer(1,8),Paragraph("※ 중복 확인 시 초과 수혜분은 환수 대상.",SM)]
pdf("장학_중복수혜_심사기준.pdf",f)

# 4 장학 서류 체크리스트(Word)
d=docx_base("장학 서류 점검 체크리스트","학생지원팀 · 장학")
dh(d,"제출 서류")
for x in ["장학신청서 (서명 포함)","소득증빙서류 (소득분위 확인)","성적증명서","통장사본","추천서(해당 시)"]:
    d.add_paragraph(x, style="List Bullet")
dh(d,"점검 포인트")
for x in ["소득분위 8분위 이하 여부","직전학기 성적 기준 충족","중복 수혜 여부(수혜내역 대조)"]:
    d.add_paragraph(x, style="List Bullet")
d.save(OUT/"장학_서류_체크리스트.docx"); print("  ✓ DOCX 장학 체크리스트")

# 5 성적이의 처리절차(Word)
d=docx_base("성적 이의신청 처리 절차","교학팀 · 학사")
dh(d,"처리 단계")
for x in ["이의신청서 접수 (성적공시 후 5일 이내)","담당 교수 확인·근거 검토","정정 사유 시 성적 정정, 아닐 시 사유 통보","처리결과 학생 통보"]:
    d.add_paragraph(x, style="List Number")
dh(d,"필요 서류")
for x in ["성적이의신청서","근거자료(과제·시험 사본 등)"]:
    d.add_paragraph(x, style="List Bullet")
d.save(OUT/"성적이의신청_처리절차.docx"); print("  ✓ DOCX 성적이의 절차")

# 6 민원 개인정보 제공 안내문(Word)
d=docx_base("개인정보 제공 관련 안내문 (학부모 성적 문의)","행정지원팀 · 민원")
d.add_paragraph("안녕하십니까. 한빛대학교 경영대학 행정지원팀입니다. 문의 주신 성적 등 학생 정보 제공과 관련하여 안내드립니다.")
dh(d,"안내 사항")
for x in ["개인정보 보호법에 따라 성인 학생의 성적은 본인 동의 없이 제3자(학부모 포함)에게 제공할 수 없습니다.",
          "정보 제공을 원하실 경우, 학생 본인의 개인정보 제공 동의서 제출이 필요합니다.",
          "동의서 제출 후 규정 범위 내에서 안내가 가능합니다."]:
    d.add_paragraph(x, style="List Bullet")
dh(d,"제공 불가 정보")
d.add_paragraph("성적, 학번, 연락처 등 개인 식별·민감정보 (동의 전)")
d.add_paragraph(""); p=d.add_paragraph(); r=p.add_run("※ 본 안내문은 synthetic_demo 입니다."); r.font.size=Pt(8.5); r.font.color.rgb=RGBColor(0x88,0x88,0x88)
d.save(OUT/"민원_개인정보제공_안내문.docx"); print("  ✓ DOCX 개인정보 안내문")

# 7 유학생 비자 서류 체크리스트(PDF)
f=hdr("유학생 비자·체류 서류 체크리스트","국제교류팀 · 유학생")
f+=[Paragraph("필수 서류",H),tbl([["서류","상태"],["비자사본(D-2)","필수"],["표준입학허가서","필수"],["재정증명서","필수"],["외국인등록증","입국 후"],["보험증서","의무"]],[80*mm,94*mm]),
    Paragraph("체류 관리",H),bl(["체류기간 만료 1개월 전 연장 신청","주소 변경 시 14일 내 신고","TOPIK 성적 요건 확인"])]
pdf("유학생_비자서류_체크리스트.pdf",f)

# 8 학사경고 상담 가이드(Word)
d=docx_base("학사경고 상담 가이드","학생성공센터 · 학생성공")
dh(d,"상담 절차")
for x in ["대상자 추출(직전학기 평점 기준)","상담 동의서 수령","1:1 상담·학습계획 수립","튜터링·심리상담 연계","후속 모니터링"]:
    d.add_paragraph(x, style="List Number")
dh(d,"필요 서류")
for x in ["상담일지","학습계획서","상담동의서(개인정보 동의 포함)"]:
    d.add_paragraph(x, style="List Bullet")
d.save(OUT/"학사경고_상담가이드.docx"); print("  ✓ DOCX 학사경고 가이드")

print("\n완료 →", OUT)
